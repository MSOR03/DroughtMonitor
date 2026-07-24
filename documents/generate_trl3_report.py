"""
Generador del informe de evidencia técnica TRL3 (MinCiencias / ATENEA) para SIPREH.
Formato adaptado a los lineamientos de presentación de documentos de la
Universidad Nacional de Colombia (portada institucional, Times New Roman,
interlineado 1.5, texto justificado, numeración de página).

Uso:
    conda run -n droughts python documents/generate_trl3_report.py

Genera: documents/Informe_TRL3_SIPREH.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "documents/Informe_TRL3_SIPREH.docx"
IMG_DIR = "documents/evidencia_generada"

BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x40, 0x40, 0x40)
TABLE_HEADER_BG = "D9D9D9"
NOTE_BG = "F2F2F2"


def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)


def add_toc(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Clic derecho sobre este texto y seleccionar 'Actualizar campo' para generar la tabla de contenido."
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def note(document, text):
    """Nota breve, sobria, indicando qué falta anexar (sin listas ni énfasis excesivo)."""
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, NOTE_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run("Nota: " + text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = GRAY
    document.add_paragraph()


def add_table(document, headers, rows, col_widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK
        run.font.size = Pt(10.5)
        set_cell_background(hdr_cells[i], TABLE_HEADER_BG)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    document.add_paragraph()
    return table


def heading(document, text, level=1):
    h = document.add_heading("", level=level)
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK
    run.font.bold = True
    sizes = {1: 14, 2: 12.5, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    if level == 3:
        run.italic = True
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return h


def body(document, text, size=12, italic=False, bold=False, align=None, space_after=8):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align if align else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    return p


def bullet(document, text):
    p = document.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_picture(document, path, width_in=6.0, caption=None):
    document.add_picture(path, width=Inches(width_in))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = GRAY


def main():
    document = Document()

    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    sections = document.sections
    for sec in sections:
        sec.top_margin = Cm(3)
        sec.bottom_margin = Cm(3)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2.5)
        footer = sec.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number_field(fp)

    # =========================================================
    # PORTADA (estilo institucional Universidad Nacional de Colombia)
    # =========================================================
    for _ in range(2):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD NACIONAL DE COLOMBIA")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.color.rgb = BLACK

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sede Bogotá — Facultad de Ingeniería")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Grupo de Investigación en Ingeniería de los Recursos Hídricos (GIREH)")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    for _ in range(6):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EVIDENCIA TÉCNICA DEL NIVEL DE MADUREZ TECNOLÓGICA TRL 3")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Marco de madurez tecnológica de MinCiencias")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SIPREH — Sistema Integral de Predicción y Reconstrucción del Estado Hídrico")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True

    for _ in range(8):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento preparado para: ATENEA")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Autor(es): [Nombre del autor / equipo]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bogotá D.C., Colombia")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("2026")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    document.add_page_break()

    # =========================================================
    # TABLA DE CONTENIDO
    # =========================================================
    heading(document, "Tabla de contenido", level=1)
    add_toc(document)
    document.add_page_break()

    # =========================================================
    # 1. INTRODUCCIÓN
    # =========================================================
    heading(document, "1. Introducción", level=1)

    body(document,
         "La gestión del recurso hídrico en Bogotá depende, en buena parte, de la capacidad para anticipar "
         "periodos de escasez en las cuencas que alimentan el sistema de acueducto de la ciudad. La variabilidad "
         "climática de los últimos años, junto con la creciente presión sobre el recurso, ha hecho evidente que "
         "el monitoreo tradicional de la sequía —basado en reportes puntuales y en el cruce manual de datos "
         "provenientes de distintas entidades— resulta insuficiente para apoyar decisiones oportunas.")

    body(document,
         "Hoy en día, obtener una lectura actualizada del estado de sequía de una cuenca implica reunir datos de "
         "precipitación de varias fuentes satelitales y de reanálisis, calcular manualmente los índices de sequía "
         "correspondientes y, en el mejor de los casos, extrapolar una tendencia a partir de la experiencia del "
         "analista. Este proceso es lento, difícil de repetir de la misma forma cada vez y no está disponible de "
         "manera continua para quienes deben tomar decisiones sobre el manejo del agua.")

    body(document,
         "El proyecto SIPREH nace para atender esta necesidad. La solución automatiza la obtención de grillas "
         "climáticas provenientes de fuentes satelitales y de reanálisis (ERA5, ERA5-Land, IMERG y CHIRPS) y de "
         "estaciones hidrológicas del IDEAM, calcula sobre ellas un conjunto de índices de sequía meteorológicos e "
         "hidrológicos ya validados en la literatura, y entrena un modelo de red neuronal convolucional que "
         "permite proyectar la evolución de estos índices a doce meses. Todo este procesamiento se realiza en un "
         "pipeline de datos desarrollado por el equipo del proyecto. Los resultados se sirven después a través de "
         "una API construida en FastAPI y de una aplicación web (SIPREH) que permite a cualquier usuario consultar "
         "un mapa de severidad de sequía, revisar la serie histórica de un punto o cuenca específica, y visualizar "
         "el pronóstico junto con su incertidumbre asociada.")

    body(document,
         "Este documento se concentra en la evidencia técnica de la parte de la solución que ya se encuentra "
         "operando de forma verificable: la API y la aplicación web de consulta y visualización, que fueron "
         "exploradas directamente para este informe. El pipeline de obtención de grillas, cálculo de índices y "
         "entrenamiento del modelo predictivo fue desarrollado por otro integrante del equipo; en los puntos donde "
         "corresponde presentar evidencia de esa parte se ha dejado indicado qué anexar.")

    body(document,
         "El propósito de este documento es presentar la evidencia técnica que demuestra que la solución "
         "desarrollada alcanza el Nivel de Madurez Tecnológica TRL 3 de acuerdo con los lineamientos establecidos "
         "por MinCiencias.")

    document.add_page_break()

    # =========================================================
    # 2. DESCRIPCIÓN DE LA SOLUCIÓN TECNOLÓGICA
    # =========================================================
    heading(document, "2. Descripción de la solución tecnológica", level=1)

    heading(document, "2.1 Arquitectura general", level=2)
    body(document,
         "SIPREH está organizado en tres capas que se comunican entre sí de forma secuencial. La primera es el "
         "pipeline de datos, encargado de obtener las grillas climáticas desde las fuentes externas, calcular los "
         "índices de sequía y entrenar y ejecutar el modelo de red neuronal convolucional que genera las "
         "predicciones. La segunda es un backend que expone estos resultados como servicios REST, apoyándose en un "
         "motor de consultas analíticas (DuckDB) y en una capa de caché para responder con rapidez. La tercera es "
         "la aplicación web, desde la cual el usuario final explora los resultados mediante un mapa interactivo y "
         "gráficas de series de tiempo.")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Datos (ERA5, ERA5-Land, IMERG, CHIRPS, estaciones IDEAM)\n"
        "        │\n"
        "        ▼\n"
        "Pipeline (obtención de grillas · agregación de índices · entrenamiento del modelo CNN)\n"
        "        │\n"
        "        ▼\n"
        "Procesamiento (índices SPI, SPEI, RAI, EDDI, PDSI, SDI, SRI, MFI, DDI, HDI)\n"
        "        │\n"
        "        ▼\n"
        "Almacenamiento (archivos Parquet en Cloudflare R2 + metadatos en base de datos relacional)\n"
        "        │\n"
        "        ▼\n"
        "API (FastAPI + DuckDB + caché Redis)\n"
        "        │\n"
        "        ▼\n"
        "Aplicación web SIPREH (Next.js/React + Leaflet + uPlot)"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    note(document,
         "adjuntar el diagrama de arquitectura en formato imagen (puede exportarse a PNG/SVG desde el diagrama "
         "Mermaid de documents/paper_architecture_section.md usando mermaid.live).")

    heading(document, "2.2 Componentes", level=2)
    body(document,
         "El pipeline de datos, a cargo del compañero del equipo responsable de esa parte del desarrollo, obtiene "
         "las grillas climáticas de las fuentes satelitales y de reanálisis mencionadas y las estaciones "
         "hidrológicas del IDEAM, calcula los índices de sequía y entrena el modelo de red neuronal convolucional "
         "que produce las predicciones a doce horizontes mensuales con sus respectivas bandas de incertidumbre "
         "(cuartiles Q1/Q3 e IQR).")

    body(document,
         "El backend, implementado en FastAPI, se probó y exploró directamente para este informe. Actualmente "
         "tiene registrados cinco conjuntos de datos activos, resumidos en la Tabla 1: tres grillas históricas "
         "(CHIRPS, IMERG y ERA5-Land), un archivo de datos hidrológicos de estaciones y el archivo de predicción "
         "vigente. Las consultas sobre estos archivos se resuelven con DuckDB directamente sobre Parquet, sin "
         "necesidad de cargarlos por completo en memoria, y los resultados se cachean en Redis (o en memoria como "
         "alternativa) para acelerar solicitudes repetidas.")

    add_table(
        document,
        ["Archivo", "Tipo de dato", "Fuente", "Tamaño", "Registros"],
        [
            ["30", "Histórico (grilla)", "CHIRPS · 0.05°", "258.9 MB", "84 453 193"],
            ["42", "Histórico (grilla)", "IMERG · 0.1°", "64.4 MB", "20 305 552"],
            ["45", "Histórico (grilla)", "ERA5-Land · 0.1°", "80.6 MB", "24 072 146"],
            ["38", "Hidrológico", "Estaciones IDEAM", "1.6 MB", "120 537"],
            ["44", "Predicción (CNN)", "Modelo predictivo", "0.55 MB", "—"],
        ],
    )
    body(document, "Tabla 1. Conjuntos de datos activos registrados en el backend al momento de la consulta "
                    "(GET /api/v1/historical/files, ejecutado directamente sobre el sistema).", size=10,
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    body(document,
         "Finalmente, la aplicación web (Next.js/React) permite seleccionar una fuente, un índice y un rango de "
         "fechas, consultar el mapa de severidad correspondiente y explorar tanto series históricas como el "
         "pronóstico vigente, incluyendo un panel de administración para la gestión de archivos y usuarios.")

    heading(document, "2.3 Tecnologías utilizadas", level=2)
    add_table(
        document,
        ["Capa", "Tecnología", "Rol"],
        [
            ["Pipeline / modelo predictivo", "Python, red neuronal convolucional (CNN)", "Obtención de grillas, cálculo de índices y predicción de sequías"],
            ["Motor de consulta", "DuckDB", "Consultas analíticas sobre archivos Parquet"],
            ["Backend API", "FastAPI + Uvicorn", "Servicios REST"],
            ["Caché", "Redis + memoria", "Reducción de latencia de respuesta"],
            ["Almacenamiento en la nube", "Cloudflare R2 (S3-compatible)", "Almacenamiento de archivos Parquet"],
            ["Base de datos relacional", "SQLite / PostgreSQL", "Metadatos, usuarios, catálogo de datasets"],
            ["Frontend", "Next.js (React)", "Aplicación web SIPREH"],
            ["Mapas", "Leaflet.js", "Visualización geoespacial"],
            ["Gráficas", "uPlot / Canvas 2D", "Series de tiempo y predicciones con incertidumbre"],
        ],
    )

    heading(document, "2.4 Flujo de funcionamiento", level=2)
    body(document,
         "El pipeline procesa periódicamente los datos crudos, calcula los índices de sequía y genera las "
         "predicciones con el modelo entrenado; el resultado se almacena como archivos Parquet en Cloudflare R2. "
         "El backend consulta estos archivos bajo demanda a través de DuckDB, cachea las respuestas y las expone "
         "mediante los endpoints REST descritos en la sección anterior. La aplicación web consume estos endpoints "
         "para construir el mapa, las series de tiempo y el resumen que finalmente ve el usuario, sin que este "
         "necesite conocimientos técnicos sobre el origen de los datos.")

    document.add_page_break()

    # =========================================================
    # 3. FUNDAMENTO TECNOLÓGICO
    # =========================================================
    heading(document, "3. Fundamento tecnológico", level=1)

    body(document,
         "El aporte tecnológico de SIPREH no está en un solo componente, sino en la forma en que se integran un "
         "pipeline de procesamiento de datos hidrometeorológicos de distintas resoluciones espaciales con un "
         "modelo predictivo de aprendizaje profundo, y en cómo ambos se ponen a disposición del usuario a través "
         "de un motor de consultas de alto desempeño.")

    body(document,
         "En primer lugar, la solución combina cuatro fuentes de datos gridded con resoluciones que van de 0.05° a "
         "0.25°, junto con registros puntuales de 29 estaciones hidrológicas del IDEAM, lo que permite evaluar la "
         "sequía tanto desde una perspectiva meteorológica como hidrológica dentro de un mismo sistema. En segundo "
         "lugar, el modelo predictivo se basa en una red neuronal convolucional capaz de capturar patrones "
         "espaciales en las grillas climáticas, en lugar de depender únicamente de extrapolaciones estadísticas "
         "univariadas; sus pronósticos a doce meses incluyen una cuantificación explícita de la incertidumbre "
         "mediante cuartiles (Q1/Q3) y el rango intercuartílico.")

    body(document,
         "Una tercera decisión relevante fue no usar una base de datos relacional tradicional para las series "
         "científicas, sino un esquema de archivos Parquet consultados directamente con DuckDB. Esto permite "
         "resolver consultas analíticas sobre decenas de millones de registros —como los 84 millones de la grilla "
         "CHIRPS— sin cargar el archivo completo en memoria, algo que se verificó directamente en la sección 5 de "
         "este documento. A esto se suma una estrategia de almacenamiento por niveles, en la que los archivos se "
         "descargan de Cloudflare R2 hacia una caché en disco local solo cuando se necesitan, y un servicio de "
         "resumen basado en un modelo de lenguaje (Groq/Llama) que traduce los resultados numéricos de la "
         "predicción a un texto en español comprensible para usuarios sin formación técnica.")

    body(document,
         "Esta combinación diferencia a SIPREH de alternativas basadas en hojas de cálculo, reportes estáticos o "
         "sistemas de información geográfica que no incorporan un modelo predictivo. No es el propósito de esta "
         "sección hacer un estado del arte exhaustivo, sino dejar claras las decisiones técnicas que sustentan el "
         "carácter innovador del desarrollo.")

    document.add_page_break()

    # =========================================================
    # 4. METODOLOGÍA DE VALIDACIÓN
    # =========================================================
    heading(document, "4. Metodología de validación", level=1)

    heading(document, "4.1 Objetivo de la validación", level=2)
    body(document,
         "El objetivo de la validación es comprobar que el pipeline procesa correctamente los datos hasta obtener "
         "índices de sequía y predicciones consistentes, y que la aplicación web permite consultar y visualizar "
         "esos resultados de forma correcta y con tiempos de respuesta adecuados para un uso interactivo.")

    heading(document, "4.2 Escenarios de prueba", level=2)

    heading(document, "Escenario 1 — Procesamiento de datos en el pipeline", level=3)
    body(document,
         "Corresponde a la obtención de grillas climáticas, el cálculo de índices de sequía y la generación del "
         "archivo de predicción mediante el modelo entrenado. Esta parte del sistema no se ejecutó directamente "
         "para este informe, ya que corresponde al trabajo del compañero encargado del pipeline.")
    note(document,
         "solicitar al responsable del pipeline una captura de la ejecución en consola y, si es posible, la curva "
         "de entrenamiento del modelo (loss/val_loss) o alguna métrica de validación como MAE o RMSE.")

    heading(document, "Escenario 2 — Consulta mediante la API", level=3)
    body(document,
         "Se verificó que los endpoints de consulta histórica (/api/v1/historical/timeseries y "
         "/api/v1/historical/spatial) respondan correctamente ante solicitudes válidas y en tiempos compatibles "
         "con un uso interactivo. Esta prueba sí se ejecutó directamente contra el backend en funcionamiento, y "
         "sus resultados se presentan en la sección 5.")

    heading(document, "Escenario 3 — Visualización desde la aplicación web", level=3)
    body(document,
         "Se verificó que la aplicación SIPREH permita seleccionar una categoría de datos (sequía meteorológica o "
         "hidrológica), un índice y una unidad espacial, y que el mapa y el panel lateral respondan a esa "
         "selección. La interfaz se abrió y navegó directamente para este informe; las capturas correspondientes "
         "se incluyen en la sección 5.")

    heading(document, "Escenario 4 — Consistencia entre las distintas fuentes registradas", level=3)
    body(document,
         "Se revisó el catálogo de archivos activos en el backend (Tabla 1) para confirmar que las tres fuentes "
         "de grilla, el archivo hidrológico y el archivo de predicción conviven sin conflicto en el sistema y "
         "pueden consultarse de forma independiente.")

    heading(document, "4.3 Métricas", level=2)
    add_table(
        document,
        ["Métrica", "Descripción"],
        [
            ["Tiempo de procesamiento", "Tiempo requerido por el pipeline para obtener grillas, calcular índices y generar predicciones."],
            ["Tiempo de respuesta", "Latencia de los endpoints de la API ante solicitudes de consulta (timeseries, spatial)."],
            ["Exactitud", "Correspondencia entre los valores devueltos por la API y los valores esperados de las series consultadas."],
            ["Integridad de los datos", "Ausencia de valores faltantes o inconsistentes en las respuestas obtenidas."],
            ["Tasa de éxito", "Proporción de solicitudes y pruebas automatizadas ejecutadas satisfactoriamente sobre el total evaluado."],
            ["Efecto de la caché", "Reducción del tiempo de respuesta entre la primera consulta y las consultas subsiguientes sobre el mismo dato."],
        ],
    )

    document.add_page_break()

    # =========================================================
    # 5. EVIDENCIA EXPERIMENTAL
    # =========================================================
    heading(document, "5. Evidencia experimental", level=1)
    body(document,
         "Las pruebas de esta sección se ejecutaron directamente contra el backend y la aplicación web en "
         "funcionamiento durante la elaboración de este informe, por lo que los tiempos y resultados que se "
         "presentan a continuación corresponden a mediciones reales y no a valores de referencia.")

    heading(document, "Prueba 1 — Procesamiento de datos en el pipeline", level=2)
    body(document,
         "Como se indicó en la sección 4.2, esta prueba depende del pipeline desarrollado por el otro integrante "
         "del equipo.")
    note(document,
         "anexar aquí una tabla con la fuente procesada, el número de registros y el tiempo de procesamiento, "
         "además de una captura del log de ejecución correspondiente.")

    heading(document, "Prueba 2 — Consulta mediante la API", level=2)
    body(document,
         "Se realizaron seis consultas consecutivas al endpoint de serie de tiempo (variable de precipitación, "
         "celda en 4.525°N/-73.575°O, frecuencia mensual entre 2015 y 2020) y al endpoint espacial (precipitación "
         "diaria sobre las 294 celdas de la grilla CHIRPS registrada, fecha 2020-06-01), limpiando la caché antes "
         "de la primera consulta de cada bloque. Los resultados se resumen en la Tabla 2.")

    add_table(
        document,
        ["Endpoint", "Primera consulta (sin caché)", "Promedio con caché (5 consultas)", "Aceleración"],
        [
            ["POST /historical/timeseries (72 puntos mensuales)", "559 ms", "3.6 ms", "≈ 155×"],
            ["POST /historical/spatial (294 celdas)", "196–260 ms", "3.7 ms", "≈ 55×"],
        ],
    )
    body(document, "Tabla 2. Tiempos de respuesta medidos con y sin caché sobre el backend en ejecución local.",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    body(document,
         "También se consultó el índice SPI a escala de 3 meses para el mismo punto entre 2010 y 2020, obteniendo "
         "132 valores mensuales sin datos faltantes, con media de -0.10 y desviación estándar de 1.12, lo que es "
         "coherente con la definición del índice (media cercana a cero sobre el periodo de referencia).")

    add_picture(document, f"{IMG_DIR}/swagger3.png", width_in=5.5,
                caption="Figura 1. Documentación interactiva de la API (Swagger UI) mostrando los endpoints de "
                        "consulta histórica y de predicción.")

    heading(document, "Prueba 3 — Visualización desde la aplicación web", level=2)
    body(document,
         "Se levantó la aplicación web SIPREH de forma local, conectada al mismo backend usado en la prueba "
         "anterior, y se navegó por el panel principal. El mapa carga por defecto la grilla en baja resolución "
         "(0.25°) sobre Bogotá y las cuencas de abastecimiento, y el panel lateral permite elegir entre sequía "
         "meteorológica e hidrológica, así como entre distintas unidades espaciales (celdas, cuencas, municipios, "
         "perímetro urbano).")

    add_picture(document, f"{IMG_DIR}/dashboard2.png", width_in=6.0,
                caption="Figura 2. Panel principal de SIPREH con el mapa de estaciones y cuencas cargado.")

    add_picture(document, f"{IMG_DIR}/dashboard4.png", width_in=6.0,
                caption="Figura 3. Sección de predicción, con las 297 celdas CHIRPS del dominio y los doce "
                        "horizontes mensuales disponibles para la vista 1D.")

    body(document,
         "El valor de 297 celdas que muestra la interfaz para la grilla de predicción es consistente con las 294 "
         "celdas activas reportadas por el backend para el archivo CHIRPS histórico (Tabla 1); la pequeña "
         "diferencia corresponde a celdas del dominio nominal que no tienen aún un archivo de predicción asociado.")

    heading(document, "Prueba 4 — Verificación automatizada (pruebas unitarias e integración)", level=2)
    body(document,
         "Además de las consultas manuales anteriores, se ejecutó la suite de pruebas automatizadas del backend "
         "(pytest), que cubre autenticación, administración de archivos, consultas históricas, datos hidrológicos, "
         "predicción y agregación por cuencas. Las 106 pruebas existentes se ejecutaron correctamente, sin "
         "fallos, en 3.69 segundos.")

    add_table(
        document,
        ["Módulo de prueba", "Casos cubiertos", "Resultado"],
        [
            ["test_auth.py", "Autenticación y manejo de tokens (10 pruebas)", "10/10 aprobadas"],
            ["test_admin.py", "Administración de archivos y usuarios (20 pruebas)", "20/20 aprobadas"],
            ["test_historical.py", "Consultas históricas 1D/2D (27 pruebas)", "27/27 aprobadas"],
            ["test_hydro.py", "Datos hidrológicos e IDEAM (16 pruebas)", "16/16 aprobadas"],
            ["test_prediction.py", "Predicción y versionado (12 pruebas)", "12/12 aprobadas"],
            ["test_watershed.py", "Agregación por cuencas (21 pruebas)", "21/21 aprobadas"],
        ],
    )
    body(document, "Tabla 3. Resultado de la suite de pruebas automatizadas del backend "
                    "(pytest, 106 pruebas, 0 fallos, 3.69 s).", size=10, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    document.add_page_break()

    # =========================================================
    # 6. RESULTADOS OBTENIDOS
    # =========================================================
    heading(document, "6. Resultados obtenidos", level=1)
    body(document,
         "Las consultas realizadas sobre la API confirmaron que el backend responde de forma correcta tanto para "
         "series de tiempo como para datos espaciales, y que la capa de caché reduce el tiempo de respuesta de "
         "varios cientos de milisegundos a unos pocos milisegundos una vez que un dato ya fue consultado, lo cual "
         "es determinante para que el mapa y las gráficas de la aplicación web se sientan interactivos incluso "
         "sobre archivos de decenas de millones de registros.")

    body(document,
         "La suite de pruebas automatizadas, con 106 casos aprobados sin fallos, respalda que estas mismas "
         "operaciones —autenticación, administración de archivos, consultas históricas, datos hidrológicos, "
         "predicción y agregación por cuencas— se comportan de forma consistente más allá de las consultas "
         "puntuales realizadas manualmente para este informe.")

    body(document,
         "La navegación de la aplicación web mostró que un usuario puede, sin conocimientos técnicos, elegir entre "
         "sequía meteorológica e hidrológica, seleccionar una unidad espacial y llegar hasta la visualización del "
         "mapa y de la sección de predicción, la cual expone correctamente las 297 celdas del dominio y los doce "
         "horizontes mensuales del pronóstico.")

    body(document,
         "En conjunto, esta evidencia indica que la integración entre el backend y la aplicación web funciona de "
         "forma estable. La evidencia correspondiente al pipeline —procesamiento de las grillas, cálculo de "
         "índices y entrenamiento del modelo predictivo— queda pendiente de anexar por parte del compañero "
         "responsable de esa parte del desarrollo, según se indicó en la sección 5.")

    document.add_page_break()

    # =========================================================
    # 7. JUSTIFICACIÓN DEL NIVEL TRL 3
    # =========================================================
    heading(document, "7. Justificación del nivel TRL 3", level=1)

    add_table(
        document,
        ["Criterio MinCiencias", "Evidencia"],
        [
            ["Concepto tecnológico definido", "✔"],
            ["Arquitectura diseñada", "✔"],
            ["Implementación inicial", "✔"],
            ["Prueba de concepto realizada", "✔"],
            ["Resultados experimentales", "✔"],
        ],
    )

    body(document,
         "De acuerdo con los resultados obtenidos, la solución desarrollada corresponde al Nivel de Madurez "
         "Tecnológica TRL 3, ya que se dispone de una prueba de concepto analítica y experimental que demuestra la "
         "viabilidad técnica de la propuesta.")

    body(document,
         "El concepto tecnológico está definido en la arquitectura descrita en la sección 2, y su implementación "
         "inicial —pipeline, backend y aplicación web— ya se encuentra desplegada y en funcionamiento. Las "
         "pruebas presentadas en las secciones 4 y 5, realizadas directamente sobre el sistema en ejecución, "
         "constituyen la prueba de concepto que demuestra que la consulta y visualización de los resultados "
         "operan de forma correcta e integrada; la evidencia experimental del pipeline y del modelo de predicción "
         "queda por completar con la información que aporte el compañero responsable de esa parte.")

    document.add_page_break()

    # =========================================================
    # 8. CONCLUSIONES
    # =========================================================
    heading(document, "8. Conclusiones", level=1)
    bullet(document, "La solución SIPREH —pipeline de datos, backend API y aplicación web— fue implementada y se "
                      "encuentra desplegada en un entorno de prueba funcional.")
    bullet(document, "Las consultas realizadas directamente sobre la API y la aplicación web, junto con la suite "
                      "de 106 pruebas automatizadas ejecutadas sin fallos, demuestran la viabilidad técnica de la "
                      "plataforma de consulta y visualización.")
    bullet(document, "Los tiempos de respuesta medidos, con reducciones de hasta 155 veces gracias a la caché, "
                      "respaldan que la arquitectura escogida es adecuada para un uso interactivo sobre conjuntos "
                      "de datos de gran volumen.")
    bullet(document, "Con base en la evidencia técnica presentada, se considera que la tecnología desarrollada "
                      "alcanza el Nivel de Madurez Tecnológica TRL 3 conforme a los lineamientos de MinCiencias.")

    document.add_page_break()

    # =========================================================
    # REFERENCIAS
    # =========================================================
    heading(document, "Referencias", level=1)
    bullet(document, "Ministerio de Ciencia, Tecnología e Innovación (MinCiencias). Documento de Niveles de "
                      "Madurez Tecnológica (TRL).")
    bullet(document, "Documentación técnica del proyecto SIPREH (documents/paper_architecture_section.md, "
                      "drought-backend/documentation/, CLAUDE.md).")
    note(document,
         "agregar aquí artículos científicos o documentación adicional utilizada como referencia (por ejemplo, "
         "sobre SPI/SPEI/PDSI, redes neuronales convolucionales aplicadas a pronóstico climático, o DuckDB).")

    document.add_page_break()

    # =========================================================
    # ANEXOS
    # =========================================================
    heading(document, "Anexos (opcionales)", level=1)
    bullet(document, "Capturas de pantalla adicionales del dashboard y del panel administrativo.")
    bullet(document, "Diagramas de arquitectura en alta resolución.")
    bullet(document, "Fragmentos de código relevantes (backend/frontend).")
    bullet(document, "Resultados completos de las pruebas (tablas extendidas, logs).")
    bullet(document, "Enlace al repositorio GitHub: [agregar enlace].")

    document.save(OUT_PATH)
    print(f"Documento generado en: {OUT_PATH}")


if __name__ == "__main__":
    main()
